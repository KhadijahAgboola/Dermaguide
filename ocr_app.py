import streamlit as st
import cv2
import openpyxl
import easyocr
import pandas as pd
from datetime import datetime
from io import BytesIO
import numpy as np
from PIL import Image
from docx import Document
import fitz

# CSS for background color
st.markdown(
    """
    <style>
        .stApp {
            background-color: #ADD8E6;
            color: black;
        }
    </style>
    """,
    unsafe_allow_html=True
)

# Title and Problem Statement
st.title("Image Text Extraction and Management System")
st.subheader("Problem Statement")
st.write("""
We need a developer to develop a tool that extracts text (OCR) from uploaded images, organizes and stores the extracted data, and provides a searchable management system. 

**Target Users**: General users and small businesses.
""")

# Sidebar Information
st.sidebar.markdown(
    """
    <style>
        .developer-box {
            background-color: #E8BCB9; /* Light pinkish-grey */
            color: black;
            padding: 15px;
            border-radius: 10px;
        }
    </style>
    """,
    unsafe_allow_html=True
)

st.sidebar.image(
    "Headshot.jpg",
    caption="Khadijat Agboola",
    width=200
)

st.sidebar.markdown(
    """
    <div class="developer-box">
        <h3>About the Developer</h3>
        <p>
            I am <strong>Khadijat Agboola</strong>, a dedicated developer with expertise in data science, 
            machine learning, and artificial intelligence. 
        </p>
        <p>
            I created this program to help simplify text extraction from images and manage the data effectively. 
            This tool extracts text from images, organizes it into structured formats like Excel, word
            and saves the data locally.
        </p>
        <p>
            This solution is still under construction, and improvements are ongoing to ensure higher accuracy 
            and additional features.
        </p>
    </div>
    """,
    unsafe_allow_html=True
)

# Choose Your Task
st.subheader("Select an action you'd like to perform with the app")
st.write("""
This tool allows you to:
1. Extract text from images (e.g., spreadsheet image) and save the result as an Excel file.
2. Extract text from scanned documents and save the result as a Word document.
""")

# User Choice
task_choice = st.radio(
    "What would you like to do?",
    ("Upload spreadsheet image or any other image", "Upload Scanned word Document")
)

def extract_text_with_paragraphs(image, reader):
    """
    Extracts text from the image, grouping lines into paragraphs based on vertical spacing.

    Parameters:
    - image: The image to process.
    - reader: EasyOCR Reader object.

    Returns:
    - A list of paragraphs as strings.
    """
    results = reader.readtext(image, detail=1, paragraph=False)
    if not results:
        return []

    # Sort results by the vertical position of bounding boxes
    results = sorted(results, key=lambda x: x[0][0][1])

    paragraphs = []
    current_paragraph = []
    previous_y = None
    line_spacing_threshold = 15  # Adjust as needed

    for (bbox, text, confidence) in results:
        top_left = bbox[0]
        y = top_left[1]  # Vertical position of the line

        if previous_y is not None and abs(y - previous_y) > line_spacing_threshold:
            # Start a new paragraph
            paragraphs.append(" ".join(current_paragraph))
            current_paragraph = []

        current_paragraph.append(text)
        previous_y = y

    # Add the last paragraph
    if current_paragraph:
        paragraphs.append(" ".join(current_paragraph))

    return paragraphs

def extract_text_from_pdf(page):
    """
    Extract text directly from a PDF page using PyMuPDF.
    """
    text = page.get_text("text")  # Get plain text
    if text.strip():  # Return text only if it exists
        return text.split('\n')  # Split into lines for better formatting
    return []



# Function to preprocess the image
def preprocess_image(image):
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    enhanced_image = cv2.adaptiveThreshold(
        gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2
    )
    return enhanced_image

# Function to extract text with improved accuracy
def extract_text(image, reader):
    preprocessed_image = preprocess_image(image)
    results = reader.readtext(preprocessed_image, detail=1, paragraph=False)
    extracted_text = [text for _, text, _ in results]
    return extracted_text

# Save data to Excel
def save_license_data(data):
    output = BytesIO()
    df = pd.DataFrame({"Extracted Text": data})
    df.to_excel(output, index=False, engine="openpyxl")
    output.seek(0)
    return output

# Save text to Word
def save_to_word(text):
    output = BytesIO()
    doc = Document()
    for line in text:
        doc.add_paragraph(line)
    doc.save(output)
    output.seek(0)
    return output

def save_to_word_with_paragraphs(paragraphs):
    """
    Save paragraphs to a Word document, preserving paragraph structure.

    Parameters:
    - paragraphs: List of paragraphs as strings.

    Returns:
    - A BytesIO object containing the Word document.
    """
    output = BytesIO()
    doc = Document()
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)
        doc.add_paragraph("")  # Add a blank line for spacing between paragraphs
    doc.save(output)
    output.seek(0)
    return output

# Function to extract text with spatial arrangement
def extract_text_with_columns(enhanced_image, reader):
    """
    Extracts text from the image along with their bounding box positions and organizes it into columns.

    Parameters:
    - image: The image to process.
    - reader: EasyOCR Reader object.

    Returns:
    - A DataFrame with the extracted text organized into columns.
    """
    results = reader.readtext(enhanced_image, detail=1, paragraph=False)
    extracted_data = []

    for (bbox, text, confidence) in results:
        top_left = bbox[0]
        bottom_right = bbox[2]
        x_center = (top_left[0] + bottom_right[0]) / 2  # Calculate x-center for column sorting
        y_center = (top_left[1] + bottom_right[1]) / 2  # Calculate y-center for row sorting
        extracted_data.append((text, x_center, y_center))

    # Sort by x-coordinate to identify columns
    extracted_data = sorted(extracted_data, key=lambda x: x[1])

    # Group text into columns based on x-coordinate proximity
    columns = []
    current_column = []
    column_threshold = 50  # Adjust as needed based on image dimensions

    for i, data in enumerate(extracted_data):
        if i == 0:
            current_column.append(data)
        else:
            # Check if the x-coordinate difference indicates a new column
            if abs(data[1] - extracted_data[i - 1][1]) > column_threshold:
                # Save the current column and start a new one
                columns.append(current_column)
                current_column = [data]
            else:
                current_column.append(data)

    # Append the last column
    if current_column:
        columns.append(current_column)

    # Sort text within each column by y-coordinate
    sorted_columns = [sorted(col, key=lambda x: x[2]) for col in columns]

    # Prepare DataFrame
    max_rows = max(len(col) for col in sorted_columns)
    data_dict = {f"Column {i + 1}": [col[j][0] if j < len(col) else "" for j in range(max_rows)] for i, col in enumerate(sorted_columns)}

    df = pd.DataFrame(data_dict)
    return df

# Updated logic for "Scan Car License Number or any other image"
if task_choice == "Upload spreadsheet image or any other image":
    uploaded_file = st.file_uploader(
        "Upload an image file of a spreadsheet (JPEG, PNG)", type=["jpg", "jpeg", "png"]
    )
    if uploaded_file:
        file_bytes = np.asarray(bytearray(uploaded_file.read()), dtype=np.uint8)
        image = cv2.imdecode(file_bytes, cv2.IMREAD_COLOR)

        # Initialize EasyOCR reader
        reader = easyocr.Reader(["en"], gpu=False)

        # Extract text organized into columns
        st.image(image, caption="Uploaded Image", use_container_width=True)
        df = extract_text_with_columns(image, reader)

        if not df.empty:
            # Save the DataFrame to an Excel file
            output_file = "columnar_extracted_text.xlsx"
            df.to_excel(output_file, index=False)

            # Display the extracted text in columns
            st.write("**Extracted Text from Scanned Image (Organized into Columns):**")
            st.write(df)

            # Provide download link for the Excel file
            with open(output_file, "rb") as file:
                st.download_button(
                    label="Download Extracted Text as Excel",
                    data=file,
                    file_name="columnar_extracted_text.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )
        else:
            st.write("No text detected in the uploaded image.")

# Process the chosen task
elif task_choice == "Upload Scanned word Document":
    uploaded_file = st.file_uploader(
        "Upload a scanned document (JPEG, PNG, PDF)", type=["jpg", "jpeg", "png", "pdf"]
    )
    if uploaded_file is not None:
        if uploaded_file.type == "application/pdf":
            pdf_data = uploaded_file.read()
            document = fitz.open(stream=pdf_data, filetype="pdf")
            st.write(f"PDF contains {len(document)} page(s). Processing pages...")

            all_paragraphs = []
            for page_number in range(len(document)):
                page = document.load_page(page_number)
                pix = page.get_pixmap()
                image = np.array(Image.open(BytesIO(pix.tobytes("png"))))

                st.image(image, caption=f"Uploaded Page {page_number + 1}", use_container_width=True)

                # Initialize EasyOCR reader
                reader = easyocr.Reader(["en"], gpu=False)
                paragraphs = extract_text_from_pdf(page)
                all_paragraphs.extend(paragraphs)

        else:
            file_bytes = np.asarray(bytearray(uploaded_file.read()), dtype=np.uint8)
            image = cv2.imdecode(file_bytes, cv2.IMREAD_COLOR)
            st.image(image, caption="Uploaded Image", use_container_width=True)
            reader = easyocr.Reader(["en"], gpu=False)
            all_paragraphs = extract_text_with_paragraphs(image, reader)

        if all_paragraphs:
            st.write("**Extracted Text (with Paragraphs):**")
            for para in all_paragraphs:
                st.write(para)
                st.write("")  # Blank line for better visual separation

            # Save to Word file
            output = save_to_word_with_paragraphs(all_paragraphs)
            st.download_button(
                label="Download Word File",
                data=output,
                file_name="extracted_text_with_paragraphs.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        else:
            st.write("No text detected in the uploaded document.")
    else:
        st.write("Please upload a scanned document to proceed.")